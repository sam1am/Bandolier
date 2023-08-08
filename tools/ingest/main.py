import os
import json
import csv
import faiss
import numpy as np
import textract
import pandas as pd
from sklearn.feature_extraction.text import TfidfVectorizer
from docx import Document
import datetime
import hashlib
import pickle
import argparse

# Define your directory here
directory = '/Users/johngarfield/Library/CloudStorage/GoogleDrive-sam.garfield@gamp.ai/Shared drives/TNano'
supported_file_types = ['docx', 'pdf', 'csv', 'txt', 'md']
documents = []
skipped_files = []
log_file = 'log_file.txt'
metadata = {}
index_name = 'tnano'  # Define your user-defined index name
workspace_directory = 'workspace'
pickle_file_path = os.path.join(workspace_directory, 'vectorizer.pickle')

parser = argparse.ArgumentParser(description='Index documents')
parser.add_argument('--reindex', action='store_true', help='Re-index documents')
parser.add_argument('--refit', action='store_true', help='Refit the model')
args = parser.parse_args()



def build_corpus(directory, supported_file_types):
    corpus = []
    for root, dirs, files in os.walk(directory):
        for file in files:
            file_path = os.path.join(root, file)
            if file.split('.')[-1] in supported_file_types:
                print("Fitting: " + file_path)
                try:
                    if file_path.endswith('.docx'):
                        doc = Document(file_path)
                        text = ' '.join([p.text for p in doc.paragraphs])
                    elif file_path.endswith('.pdf'):
                        text = textract.process(file_path, method='pdfminer').decode('utf-8')
                    elif file_path.endswith('.csv'):
                        df = pd.read_csv(file_path)
                        text = ' '.join(df.columns) + ' ' + ' '.join(df.values.flatten().astype(str))
                    else: # For txt and md files
                        text = textract.process(file_path).decode('utf-8')
                    corpus.append(text)
                except Exception as e:
                    print(f"Error fitting {file_path}: {str(e)}")
    return corpus

def refit_model(directory, supported_file_types, pickle_file_path):
    # Build the corpus
    corpus = build_corpus(directory, supported_file_types)

    # Fit the vectorizer on the corpus
    vectorizer = TfidfVectorizer()
    vectorizer.fit(corpus)
    print("\n\nDumping the vectorizer to disk...\n\n")
    
    # Save the fitted vectorizer to a pickle file
    pickle.dump(vectorizer, open(pickle_file_path, "wb"))
    
    return vectorizer


# Check if the pickle file exists
if not os.path.exists(pickle_file_path):
    
    # Build the corpus
    corpus = build_corpus(directory, supported_file_types)

    # Fit the vectorizer on the corpus
    vectorizer = TfidfVectorizer()
    vectorizer.fit(corpus)
    print("\n\nDumping the vectorizer to disk...\n\n")
    # Save the fitted vectorizer to a pickle file
    pickle.dump(vectorizer, open(pickle_file_path, "wb"))
else:
    # Load the fitted vectorizer from the pickle file
    print("\n\nLoading the vectorizer from disk...\n\n")
    vectorizer = pickle.load(open(pickle_file_path, "rb"))

if args.refit:
    vectorizer = refit_model(directory, supported_file_types, pickle_file_path)
else:
    # Load the fitted vectorizer from the pickle file
    print("\n\nLoading the vectorizer from disk...\n\n")
    vectorizer = pickle.load(open(pickle_file_path, "rb"))

num_features = len(vectorizer.get_feature_names_out())
dimension = num_features  # The dimension of your vectors
number_of_vectors = len(metadata)  # The number of vectors

# Define the full path for the index and the metadata file
index_path = os.path.join(workspace_directory, index_name)
index_file_path = os.path.join(index_path, f'{index_name}.index')
metadata_file_path = os.path.join(index_path, 'metadata.json')

# Make sure the directory for the index exists
os.makedirs(index_path, exist_ok=True)

# Check if the index file exists or if full re-indexing is requested
if args.reindex or not os.path.exists(index_file_path):
    # Delete the existing index and metadata files
    if os.path.exists(index_file_path):
        os.remove(index_file_path)
    if os.path.exists(metadata_file_path):
        os.remove(metadata_file_path)
    # Create the base index
    index_base = faiss.IndexFlatL2(dimension)
    # Create the IDMap index
    index = faiss.IndexIDMap(index_base)
    metadata = {}
else:
    # Load the existing index and metadata
    index = faiss.read_index(index_file_path)
    with open(metadata_file_path, 'r') as f:
        metadata = json.load(f)


for root, dirs, files in os.walk(directory):
    for file in files:
        file_path = os.path.join(root, file)
        if file.split('.')[-1] in supported_file_types:
            # Generate a unique ID for each file based on its content
            print("Checking: " + file_path)
            with open(file_path, 'rb') as file:
                file_id = hashlib.md5(file.read()).hexdigest()

            if file_id not in metadata:
                documents.append(file_path)
            else:
                print(f"Skipping {file_path} - already indexed.")
        else:
            skipped_files.append(file_path)

# Log the skipped files
with open(log_file, 'w') as f:
    for skipped_file in skipped_files:
        f.write(skipped_file + '\n')

# Read each document, extract its text, and vectorize it
for document in documents:
    print("Indexing: " + document)
    try:
        # Extract text from the document based on its file type
        if document.endswith('.docx'):
            doc = Document(document)
            text = ' '.join([p.text for p in doc.paragraphs])
        elif document.endswith('.pdf'):
            # For pdf files
            text = textract.process(document, method='pdfminer').decode('utf-8')
        elif document.endswith('.csv'):
            df = pd.read_csv(document)
            text = ' '.join(df.columns) + ' ' + ' '.join(df.values.flatten().astype(str))
        else: # For txt and md files
            text = textract.process(document).decode('utf-8')

        # Vectorize the text and add it to the index
        vector = vectorizer.transform([text]).toarray()[0]
        index.add_with_ids(np.array([vector], dtype=np.float32), np.array([len(metadata)]))
        # Generate a unique ID for each file based on its content
        file_id = hashlib.md5(open(document, 'rb').read()).hexdigest()
        index_id = len(metadata)
        metadata[file_id] = {
            'file_path': document,
            'file_type': document.split('.')[-1],
            'indexing_date': str(datetime.datetime.now()),
            'index_id': index_id,
            # Add any other metadata here...
        }
        index.add_with_ids(np.array([vector], dtype=np.float32), np.array([index_id]))

    except Exception as e:
        print(f"Error processing {document}: {str(e)}")
        skipped_files.append(document)

# Save the index to a file
faiss.write_index(index, index_file_path)


number_of_vectors = index.ntotal
# Compute the memory usage in bytes
memory_usage_bytes = 4 * dimension * number_of_vectors
# Convert to more human-readable units
memory_usage_kilobytes = memory_usage_bytes / 1024
memory_usage_megabytes = memory_usage_kilobytes / 1024
memory_usage_gigabytes = memory_usage_megabytes / 1024

# Print the memory usage
print(f"Estimated memory usage of the index: {memory_usage_bytes} bytes, "
      f"{memory_usage_kilobytes} KB, {memory_usage_megabytes} MB, {memory_usage_gigabytes} GB")

# Save the metadata to a file
with open(metadata_file_path, 'w') as f:
    json.dump(metadata, f)